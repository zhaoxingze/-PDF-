import unittest
import uuid
from pathlib import Path

import fitz
from docx import Document

from src.core.docx_reader import DocxReader
from src.core.docx_writer import DocxWriter
from src.core.pdf_reader import PDFReader
from src.core.pdf_writer import PDFWriter
from src.core.translator import GoogleTranslator


TMP_ROOT = Path(__file__).resolve().parents[1] / ".test_tmp"
TMP_ROOT.mkdir(exist_ok=True)


def test_path(name: str) -> Path:
    return TMP_ROOT / f"{uuid.uuid4().hex}_{name}"


class DummyGoogleTranslator(GoogleTranslator):
    def __init__(self):
        super().__init__()
        self.calls = []

    def translate_single(self, text: str) -> str:
        self.calls.append(text)
        return f"zh:{text}"


class TranslationPipelineTests(unittest.TestCase):
    def test_google_batch_skips_non_english_text_and_reuses_duplicate_results(self):
        translator = DummyGoogleTranslator()

        result = translator.translate_batch([
            "Hello world",
            "Hello world",
            "12345",
            "\u4e2d\u6587\u5185\u5bb9",
            "A",
        ])

        self.assertEqual(result, ["zh:Hello world", "zh:Hello world", "", "", ""])
        self.assertEqual(translator.calls, ["Hello world"])

    def test_docx_reader_and_writer_keep_body_and_table_order(self):
        input_path = test_path("reader_writer_input.docx")
        output_path = test_path("reader_writer_output.docx")

        doc = Document()
        doc.add_paragraph("First paragraph")
        table = doc.add_table(rows=1, cols=1)
        table.cell(0, 0).text = "Table paragraph"
        doc.add_paragraph("Last paragraph")
        doc.save(input_path)

        paragraphs = DocxReader(str(input_path)).extract_paragraphs()

        self.assertEqual(
            [p.text for p in paragraphs],
            ["First paragraph", "Table paragraph", "Last paragraph"],
        )

        writer = DocxWriter(str(input_path))
        writer.replace_paragraphs(paragraphs, ["zh:first", "zh:table", "zh:last"])
        writer.save(str(output_path))

        out = Document(output_path)
        self.assertEqual(out.paragraphs[0].text, "zh:first")
        self.assertEqual(out.tables[0].cell(0, 0).paragraphs[0].text, "zh:table")
        self.assertEqual(out.paragraphs[1].text, "zh:last")

    def test_pdf_reader_drops_exact_overlapping_duplicate_text(self):
        pdf_path = test_path("duplicate.pdf")
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((72, 72), "Repeated text", fontsize=12)
        page.insert_text((72, 72), "Repeated text", fontsize=12)
        page.insert_text((72, 96), "Next text", fontsize=12)
        doc.save(pdf_path)
        doc.close()

        reader = PDFReader(str(pdf_path))
        try:
            texts = [block.merged_text for block in reader.extract_page_blocks(0)]
        finally:
            reader.close()

        self.assertEqual(texts, ["Repeated text", "Next text"])

    def test_pdf_writer_embeds_readable_chinese_translation(self):
        input_path = test_path("font_input.pdf")
        output_path = test_path("font_output.pdf")
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((72, 72), "Original text", fontsize=12)
        doc.save(input_path)
        doc.close()

        reader = PDFReader(str(input_path))
        try:
            blocks = reader.extract_page_blocks(0)
            font_path = Path(__file__).resolve().parents[1] / "fonts" / "NotoSansSC-Regular.ttf"
            writer = PDFWriter(reader.doc, str(font_path))
            writer.replace_text_on_page(0, blocks, ["\u4e2d\u6587\u8bd1\u6587"])
            writer.save(str(output_path))
        finally:
            reader.close()

        out = fitz.open(output_path)
        try:
            self.assertIn("\u4e2d\u6587\u8bd1\u6587", out[0].get_text())
        finally:
            out.close()


if __name__ == "__main__":
    unittest.main()
